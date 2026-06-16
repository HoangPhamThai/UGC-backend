import copy
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table, _Row

from app.modules.reports.domain.errors import ReportValidationError

REQUIRED_TEMPLATE_TOKENS = ("{creator_name}", "{final_award}", "{article_platform}")
TEMPLATE_PATH = str(Path(__file__).parent / "templates" / "acceptance_report.docx")

_ROW_TOKEN = "{article_platform}"
_ROW_KEYS = (
    "article_platform",
    "article_id_autoinc",
    "article_on_air",
    "article_link",
    "article_view",
    "article_image",
    "article_bonus_money",
)


def _replace_in_paragraph(paragraph, mapping: dict) -> None:
    full = "".join(run.text for run in paragraph.runs)
    if "{" not in full:
        return
    new = full
    for key, value in mapping.items():
        new = new.replace("{" + key + "}", value)
    if new == full:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""


def _iter_cell_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph


def _row_contains_token(row, token: str) -> bool:
    return any(token in cell.text for cell in row.cells)


def _insert_image_in_cell(cell, image_bytes: bytes) -> None:
    """Clear cell content and insert image, scaled to fit cell width."""
    max_width = cell.width or 1_828_800  # fallback: 2 inches in EMU
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    pic = run.add_picture(BytesIO(image_bytes))
    if pic.width > max_width:
        scale = max_width / pic.width
        pic.width = max_width
        pic.height = int(pic.height * scale)


_ZWSP = "​"  # ZERO WIDTH SPACE — lets Word break long URLs in fixed-width cells
_LINK_BREAK_AFTER = "/.?&-_="  # chars after which a ZWSP is inserted


def _add_break_opportunities(url: str) -> str:
    """Insert a zero-width space after URL separator characters so Word can wrap
    the link inside a fixed-width cell. ZWSP is invisible; copy-paste stays intact."""
    if not url:
        return url
    out = []
    for ch in url:
        out.append(ch)
        if ch in _LINK_BREAK_AFTER:
            out.append(_ZWSP)
    return "".join(out)


def _fill_row(row, item: dict) -> None:
    image_bytes: Optional[bytes] = item.get("_image_bytes")
    mapping = {k: str(item.get(k, "")) for k in _ROW_KEYS if k != "article_image"}
    mapping["article_link"] = _add_break_opportunities(mapping["article_link"])
    for cell in row.cells:
        cell_text = " ".join(p.text for p in cell.paragraphs)
        if "{article_image}" in cell_text:
            if image_bytes:
                _insert_image_in_cell(cell, image_bytes)
            else:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, {"article_image": ""})
        else:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, mapping)


def render_acceptance_report(
    *,
    scalars: dict,
    line_items: list[dict],
    template_bytes: Optional[bytes] = None,
    line_item_images: Optional[dict[str, bytes]] = None,
) -> bytes:
    """Render acceptance report .docx. `line_item_images` maps article_id to image bytes;
    when provided, images are embedded directly into the {article_image} table cell."""
    document = Document(BytesIO(template_bytes)) if template_bytes else Document(TEMPLATE_PATH)
    scalar_map = {k: str(v) for k, v in scalars.items()}

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, scalar_map)

    tables = list(document.tables)
    article_table_idx = None
    template_row = None
    for idx, table in enumerate(tables):
        for row in table.rows:
            if _row_contains_token(row, _ROW_TOKEN):
                article_table_idx = idx
                template_row = row
                break
        if article_table_idx is not None:
            break

    for idx, table in enumerate(tables):
        for row in table.rows:
            if idx == article_table_idx and template_row is not None and row._tr is template_row._tr:
                continue
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, scalar_map)

    if template_row is not None and article_table_idx is not None:
        article_table = tables[article_table_idx]
        anchor_tr = template_row._tr
        ref = anchor_tr
        images = line_item_images or {}
        for item in line_items:
            enriched = {**item, "_image_bytes": images.get(item.get("article_id", ""))}
            new_tr = copy.deepcopy(anchor_tr)
            ref.addnext(new_tr)
            ref = new_tr
            _fill_row(_Row(new_tr, article_table), enriched)
        anchor_tr.getparent().remove(anchor_tr)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def validate_template_bytes(data: bytes) -> None:
    """Raise ReportValidationError if `data` is not a usable acceptance-report
    template (not a valid .docx, or missing required tokens)."""
    try:
        document = Document(BytesIO(data))
    except Exception as exc:  # noqa: BLE001 — any open failure means invalid upload
        raise ReportValidationError("File không phải .docx hợp lệ") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    missing = [t for t in REQUIRED_TEMPLATE_TOKENS if t not in text]
    if missing:
        raise ReportValidationError(f"Template thiếu token bắt buộc: {', '.join(missing)}")
