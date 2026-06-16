from io import BytesIO
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")  # skip if python-docx not installed locally

from app.modules.reports.rendering import TEMPLATE_PATH, render_acceptance_report, validate_template_bytes, _ZWSP as ZWSP, _CELL_SAFETY_MARGIN_EMU
from app.modules.reports.domain.errors import ReportValidationError


def _all_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _scalars() -> dict:
    return {
        "created_at": "14/06/2026",
        "creator_name": "Nguyen Van A",
        "creator_date_of_birth": "1990-01-15",
        "creator_social_id": "0123",
        "creator_social_id_date_of_issue": "2015-01-01",
        "creator_social_id_place_of_issue": "HCMC",
        "creator_primary_address": "Addr 1",
        "creator_current_address": "Addr 2",
        "creator_tax_number": "TAX1",
        "creator_bank_account_number": "ACC1",
        "creator_bank": "Bank",
        "creator_bank_branch": "Branch",
        "total_approved_articles": "2",
        "total_articles": "2",
        "article_award_price": "500000",
        "total_award": "1000000",
        "tax": "100000",
        "final_award": "900000",
        "final_award_verbal": "Chín trăm nghìn",
    }


def _items() -> list[dict]:
    return [
        {
            "article_id": "art_1",
            "article_platform": "tiktok", "article_id_autoinc": "1",
            "article_on_air": "2026-06-01", "article_link": "https://tt/1",
            "article_view": "100", "article_image": "", "article_bonus_money": "  ",
        },
        {
            "article_id": "art_2",
            "article_platform": "threads", "article_id_autoinc": "2",
            "article_on_air": "2026-06-02", "article_link": "https://th/2",
            "article_view": "200", "article_image": "", "article_bonus_money": "  ",
        },
    ]


def test_template_exists():
    assert Path(TEMPLATE_PATH).exists()


def test_render_substitutes_scalars_and_clones_rows():
    out = render_acceptance_report(scalars=_scalars(), line_items=_items())
    document = docx.Document(BytesIO(out))
    text = _all_text(document)

    text_no_zwsp = text.replace(ZWSP, "")
    assert "{" not in text and "}" not in text
    assert "Nguyen Van A" in text
    assert "Chín trăm nghìn" in text
    assert "https://tt/1" in text_no_zwsp and "https://th/2" in text_no_zwsp
    assert "tiktok" in text and "threads" in text


def test_render_handles_single_item():
    out = render_acceptance_report(scalars=_scalars(), line_items=_items()[:1])
    document = docx.Document(BytesIO(out))
    text = _all_text(document)
    text_no_zwsp = text.replace(ZWSP, "")
    assert "https://tt/1" in text_no_zwsp and "https://th/2" not in text_no_zwsp
    assert "{article_platform}" not in text


def test_validate_template_accepts_the_vendored_default():
    data = Path(TEMPLATE_PATH).read_bytes()
    validate_template_bytes(data)  # must not raise


def test_validate_template_rejects_non_docx():
    with pytest.raises(ReportValidationError):
        validate_template_bytes(b"not a docx")


def test_render_uses_provided_template_bytes():
    data = Path(TEMPLATE_PATH).read_bytes()
    out = render_acceptance_report(scalars=_scalars(), line_items=_items(), template_bytes=data)
    document = docx.Document(BytesIO(out))
    text = _all_text(document)
    assert "Nguyen Van A" in text and "{" not in text


def _minimal_docx_with_image_token() -> bytes:
    """In-memory .docx with one table containing {article_image} in col 2."""
    from docx import Document as D
    doc = D()
    tbl = doc.add_table(rows=1, cols=3)
    tbl.cell(0, 0).text = "{article_platform}"
    tbl.cell(0, 1).text = "{article_id_autoinc}"
    tbl.cell(0, 2).text = "{article_image}"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_png(w: int, h: int) -> bytes:
    """Solid red w×h PNG (no pHYs chunk), no external deps."""
    import struct, zlib
    def chunk(t, d):
        return struct.pack('>I', len(d)) + t + d + struct.pack('>I', zlib.crc32(t + d) & 0xFFFFFFFF)
    ihdr = struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * w for _ in range(h))
    idat = zlib.compress(raw)
    return b'\x89PNG\r\n\x1a\n' + chunk(b'IHDR', ihdr) + chunk(b'IDAT', idat) + chunk(b'IEND', b'')


def _1x1_png() -> bytes:
    return _make_png(1, 1)


def test_render_embeds_image_replaces_token():
    tpl = _minimal_docx_with_image_token()
    png = _1x1_png()
    item = {
        "article_id": "art_1", "article_platform": "tiktok",
        "article_id_autoinc": "1", "article_on_air": "2026-06-01",
        "article_link": "https://x", "article_view": "100",
        "article_image": "k/art_1.png", "article_bonus_money": "  ",
    }
    out = render_acceptance_report(
        scalars={}, line_items=[item],
        template_bytes=tpl,
        line_item_images={"art_1": png},
    )
    doc = docx.Document(BytesIO(out))
    assert "{article_image}" not in doc.tables[0].cell(0, 2).text


def test_render_clears_token_when_no_image():
    tpl = _minimal_docx_with_image_token()
    item = {
        "article_id": "art_1", "article_platform": "tiktok",
        "article_id_autoinc": "1", "article_on_air": "2026-06-01",
        "article_link": "https://x", "article_view": "100",
        "article_image": "", "article_bonus_money": "  ",
    }
    out = render_acceptance_report(scalars={}, line_items=[item], template_bytes=tpl)
    doc = docx.Document(BytesIO(out))
    assert "{article_image}" not in doc.tables[0].cell(0, 2).text


def _cell_text(document, table_idx, row_idx, col_idx):
    return document.tables[table_idx].rows[row_idx].cells[col_idx].text


def test_link_gets_break_opportunities():
    # link col = index 3, data row = index 1 (after header) in the article table
    link = "https://www.youtube.com/watch?v=abc-def_123&x=1"
    item = {
        "article_id": "art_1", "article_platform": "tiktok",
        "article_id_autoinc": "1", "article_on_air": "2026-06-01",
        "article_link": link, "article_view": "100",
        "article_image": "", "article_bonus_money": "  ",
    }
    out = render_acceptance_report(scalars=_scalars(), line_items=[item])
    document = docx.Document(BytesIO(out))
    cell_text = _cell_text(document, 0, 1, 3)

    assert ZWSP in cell_text  # break opportunities inserted
    assert cell_text.replace(ZWSP, "") == link  # stripping ZWSP yields the intact URL
    # ZWSP sits right after URL separator characters
    assert "/" + ZWSP in cell_text
    assert "?" + ZWSP in cell_text
    assert "-" + ZWSP in cell_text


def _solid_png(w: int, h: int) -> bytes:
    return _make_png(w, h)


def _grid_col_emu(document, table_idx, col_idx):
    from docx.oxml.ns import qn
    from docx.shared import Twips
    grid = document.tables[table_idx]._tbl.find(qn('w:tblGrid'))
    cols = grid.findall(qn('w:gridCol'))
    return int(Twips(int(float(cols[col_idx].get(qn('w:w'))))))


def _image_extent_cx(document, table_idx, row_idx, col_idx):
    from docx.oxml.ns import qn
    cell = document.tables[table_idx].rows[row_idx].cells[col_idx]
    ext = cell._tc.findall('.//' + qn('wp:extent'))
    return int(ext[0].get('cx')) if ext else None


def _native_picture_width(image_bytes):
    """Width (EMU) python-docx assigns when embedding the image at native size."""
    from docx import Document as D
    run = D().add_paragraph().add_run()
    return run.add_picture(BytesIO(image_bytes)).width


def _image_item(png_key="k/art_1.png"):
    return {
        "article_id": "art_1", "article_platform": "tiktok",
        "article_id_autoinc": "1", "article_on_air": "2026-06-01",
        "article_link": "https://x/1", "article_view": "100",
        "article_image": png_key, "article_bonus_money": "  ",
    }


def test_large_image_scaled_to_column_width():
    # image col = index 5 in the article table; data row = index 1
    big = _solid_png(1000, 500)
    out = render_acceptance_report(
        scalars=_scalars(), line_items=[_image_item()],
        line_item_images={"art_1": big},
    )
    document = docx.Document(BytesIO(out))
    cx = _image_extent_cx(document, 0, 1, 5)
    col_emu = _grid_col_emu(document, 0, 5)

    assert cx is not None
    assert cx != 1_828_800            # not the old 2-inch fallback (the bug)
    assert cx == col_emu - _CELL_SAFETY_MARGIN_EMU  # capped to column minus padding


def test_small_image_not_upscaled():
    small = _1x1_png()  # smaller than the column → must be left at native size
    native = _native_picture_width(small)
    out = render_acceptance_report(
        scalars=_scalars(), line_items=[_image_item()],
        line_item_images={"art_1": small},
    )
    document = docx.Document(BytesIO(out))
    cx = _image_extent_cx(document, 0, 1, 5)
    col_emu = _grid_col_emu(document, 0, 5)
    assert native < col_emu   # precondition: it does fit without scaling
    assert cx == native       # not upscaled to column width, not shrunk
